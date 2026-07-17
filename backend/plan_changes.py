"""
One-off script to generate PlanChanges.docx — a summary of what sections
in the planned Week-5 architecture need rewriting to match the current TrustAudit
codebase.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── Title ──
    title = doc.add_heading("TrustAudit — Plan Amendments", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Changes required in the planned Week-5 architecture document "
        "to match the current TrustAudit implementation."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: 2026-07-16").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Helper ──
    def h(text, level=1):
        return doc.add_heading(text, level=level)

    def body(text):
        return doc.add_paragraph(text)

    def bold(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return p

    def table(headers, rows):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, hdr in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = hdr
            from docx.oxml.ns import qn
            shading = cell._tc.get_or_add_tcPr()
            shd = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "E1F5EE",
            })
            shading.append(shd)
            for paragraph in cell.paragraphs:
                for run_ in paragraph.runs:
                    run_.bold = True
        for row in rows:
            for i, val in enumerate(row):
                tbl.add_row().cells[i].text = val
        doc.add_paragraph()
        return tbl

    # ══════════════════════════════════════════════════════════════
    # 1. DOCUMENT INGESTION LAYER
    # ══════════════════════════════════════════════════════════════
    h("1. Document Ingestion Layer")
    body(
        "The planned document states support for PDF, JPG, PNG, and scanned documents. "
        "Our current implementation additionally supports TIFF, TIF, and BMP (see "
        "SUPPORTED_EXTENSIONS in app/graph.py)."
    )
    bold("Required change:")
    body(
        "Update the supported format list to include TIFF and BMP. Mention that file "
        "discovery is driven by os.scandir() with a frozenset of extensions. Also note "
        "that no file-integrity validation is performed at ingestion time (acceptable "
        "for a Week-5 prototype)."
    )

    # ══════════════════════════════════════════════════════════════
    # 2. DOCUMENT RENDERING
    # ══════════════════════════════════════════════════════════════
    h("2. Document Rendering")
    body(
        "Already accurate. The planned section describes PDF → PyMuPDF → PIL, all "
        "in-memory without temp files. This is identical to our _pdf_to_images() method "
        "in app/vlm.py. No changes needed."
    )

    # ══════════════════════════════════════════════════════════════
    # 3. OCR AND LAYOUT EXTRACTION  ← MAJOR REWRITE
    # ══════════════════════════════════════════════════════════════
    h("3. OCR and Layout Extraction  —  MAJOR REWRITE")
    body(
        "The planned section describes Tesseract OCR with word-level bounding boxes, "
        "confidence scores, and coordinate extraction. Our project does NOT use Tesseract "
        "or any dedicated OCR engine. Instead, document page images are sent directly to "
        "Gemini 2.5 Flash as base64-encoded JPEGs via the langchain-google-genai library. "
        "The VLM performs native vision-based text extraction — there is no separate OCR step."
    )
    bold("Rewrite to:")
    body(
        "\"3. Vision-Language Extraction\n\n"
        "Once rendered, page images are converted to base64 JPEG strings and embedded in a "
        "multipart Gemini message alongside a structured prompt. The model (gemini-2.5-flash) "
        "performs end-to-end text extraction, field identification, and spatial reasoning from "
        "the raw pixel data. No Tesseract, no explicit bounding-box logic, and no confidence "
        "scores are computed.\n\n"
        "This approach trades offline capability for significantly higher accuracy on "
        "challenging layouts, handwriting, and low-quality scans. API availability is a "
        "prerequisite.\""
    )

    # ══════════════════════════════════════════════════════════════
    # 4. SIMULATION / FALLBACK ENGINE
    # ══════════════════════════════════════════════════════════════
    h("4. Simulation / Fallback Engine  —  DELETE OR MOVE")
    body(
        "The planned section describes predefined coordinate templates for offline field "
        "extraction when OCR is unavailable. Our project has no such fallback mechanism. "
        "If the Gemini API call fails, the system returns zero-score defaults and logs "
        "the error."
    )
    bold("Required change:")
    body(
        "Delete this section from the Week-5 scope, or move it to a \"Future Work\" "
        "section with a note that offline template-based extraction is deferred to "
        "post-Week 5."
    )

    # ══════════════════════════════════════════════════════════════
    # 5. SPATIAL HEURISTIC ENGINE  ← MAJOR REWRITE
    # ══════════════════════════════════════════════════════════════
    h("5. Spatial Heuristic Engine  —  MAJOR REWRITE")
    body(
        "The planned section describes explicit Euclidean-distance keyword proximity logic "
        "(e.g., find 'TOTAL' label, search nearby bounding boxes for the closest number). "
        "Our project does not implement any explicit spatial heuristics. Spatial reasoning "
        "is handled implicitly by the VLM when it receives full-page images."
    )
    bold("Rewrite to:")
    body(
        "\"5. Vision-Based Spatial Reasoning\n\n"
        "Spatial field localization is delegated to the VLM. The model receives complete "
        "page images and uses its native visual understanding to associate labels with their "
        "values (e.g., finding the monetary figure adjacent to the 'TOTAL' label). No "
        "explicit coordinate matching, Euclidean distance, or bounding-box overlap logic is "
        "implemented. This approach is less transparent but more robust to layout variations.\n\n"
        "Future work may add explicit bounding-box extraction from the VLM response for "
        "visual overlay and debugging.\""
    )

    # ══════════════════════════════════════════════════════════════
    # 6. SEMANTIC PARSING LAYER  ← REWRITE
    # ══════════════════════════════════════════════════════════════
    h("6. Semantic Parsing Layer  —  REWRITE")
    body(
        "The planned section describes post-processing OCR output into structured fields. "
        "Our project achieves this in a single VLM call during the summarize step."
    )
    bold("Rewrite to:")
    body(
        "\"6. Structured Field Extraction\n\n"
        "During Step 2 (summarize_documents), Gemini is prompted to return a JSON object "
        "with fields including:\n"
        "  - document_type (invoice, purchase_order, receipt, etc.)\n"
        "  - summary (2–4 sentence natural-language description)\n"
        "  - language\n"
        "  - key_fields (document_number, date, total_amount, currency, vendor_name, "
        "customer_name)\n\n"
        "The VLM response is parsed with json.loads() using a deduplication hook for "
        "resilience. If parsing fails, fallback defaults are used.\""
    )

    # ══════════════════════════════════════════════════════════════
    # 7. COMPLIANCE RULES ENGINE
    # ══════════════════════════════════════════════════════════════
    h("7. Compliance Rules Engine")
    body(
        "The concept is correct, but the planned example rules should be replaced with "
        "our actual 12-rule checklist (R001–R012) loaded from checklist.md at runtime. "
        "Rules include severity (critical/high/medium/low) and a mandatory flag."
    )
    bold("Required change:")
    body(
        "Replace the placeholder example rules with the full checklist from checklist.md. "
        "Note that rules are parsed from Markdown using regex and structured into "
        "AuditRule Pydantic models."
    )

    # ══════════════════════════════════════════════════════════════
    # 8. COMPLIANCE SCORING  ← REWRITE
    # ══════════════════════════════════════════════════════════════
    h("8. Compliance Scoring  —  REWRITE")
    body(
        "The planned section describes weighted scoring (GST=35, Invoice#=25, Date=20, "
        "Total=20). Our project uses equal-weight percentage scoring."
    )
    bold("Rewrite to:")
    body(
        "\"8. Compliance Scoring\n\n"
        "All rules carry equal weight. A document's score is:\n"
        "  score = (rules_in_compliance / total_rules) × 100\n\n"
        "The 'passed' boolean requires ALL mandatory rules to be in compliance. "
        "Optional rules contribute to the percentage score but do not alone cause a FAIL.\n\n"
        "Overall score = mean of all per-document scores. "
        "A 95% prediction interval is computed from per-document score variance.\n\n"
        "Weighted scoring (with configurable rule weights) is deferred to a future iteration.\""
    )

    # ══════════════════════════════════════════════════════════════
    # 9. RISK CLASSIFICATION
    # ══════════════════════════════════════════════════════════════
    h("9. Risk Classification")
    body(
        "Already accurate. Our thresholds are identical: ≥80 PASS, 50–79 REVIEW REQUIRED, "
        "<50 FAIL. No changes needed."
    )

    # ══════════════════════════════════════════════════════════════
    # 10. AUDIT TRAIL GENERATION
    # ══════════════════════════════════════════════════════════════
    h("10. Audit Trail Generation")
    body(
        "Our FailedChecklistItem model stores rule_id, rule_title, description, severity, "
        "evidence, and page_number — matching the plan's requirements. Additionally, "
        "evidence text is cross-referenced against the original checklist via "
        "_safe_failed_item() to ensure rule metadata is authoratitive even if the VLM "
        "returns incorrect titles or severities."
    )
    bold("Required change:")
    body(
        "Add a note about the cross-reference safeguard. Otherwise accurate — no "
        "rewrite needed."
    )

    # ══════════════════════════════════════════════════════════════
    # 11. REPORT GENERATION  ← REWRITE
    # ══════════════════════════════════════════════════════════════
    h("11. Report Generation  —  REWRITE")
    body(
        "The planned section describes PDF generation via ReportLab. Our project generates "
        "a .docx file via python-docx."
    )
    bold("Rewrite to:")
    body(
        "\"11. Report Generation\n\n"
        "After the audit completes, a .docx report is generated containing:\n"
        "  - Title page with audit name, timestamp, and elapsed time\n"
        "  - Overall result (PASSED / REVIEW REQUIRED / FAILED) with score\n"
        "  - Prediction interval (95% confidence range)\n"
        "  - Document count breakdown\n"
        "  - Per-document sections, each with:\n"
        "      * Embedded JPEG thumbnail preview (first page)\n"
        "      * Score and pass/fail status\n"
        "      * Remarks\n"
        "      * 3-column failed-rules table (Rule, Severity, Evidence)\n\n"
        "The report is served as a downloadable file via FastAPI's Response with "
        "Content-Disposition: attachment.\"\n\n"
        "Note: PDF export via ReportLab is deferred."
    )

    # ══════════════════════════════════════════════════════════════
    # 12. USER INTERFACE  ← FULL REWRITE
    # ══════════════════════════════════════════════════════════════
    h("12. User Interface  —  FULL REWRITE")
    body(
        "The planned section describes a Streamlit interface with file upload, "
        "bounding-box overlays, confidence tables, and PDF export. Our project uses a "
        "completely different tech stack."
    )
    bold("Rewrite to:")
    body(
        "\"12. User Interface\n\n"
        "Technology: React 19 + TypeScript + Vite + Tailwind CSS v4\n\n"
        "Layout (single page, max-width 960px):\n\n"
        "  a. Folder Input — text field for documents directory path, with 'Run audit' "
        "button (teal outline) and conditional 'Download .docx' button (blue outline).\n\n"
        "  b. Processing Animation — magnifying-glass SVG orbiting on a dashed circular "
        "track with a horizontal scan-line oscillating left-to-right, displayed while the "
        "pipeline runs.\n\n"
        "  c. Audit Summary — overall result badge, score with 95% prediction interval, "
        "document counts (processed/passed/failed), and natural-language summary.\n\n"
        "  d. Summary Metrics — 4-column grid: Documents Processed | Passed | Failed | "
        "Overall Score.\n\n"
        "  e. Document Cards — one per document with:\n"
        "      * 120×120 JPEG thumbnail on the left\n"
        "      * Document name (truncated), score, pass/fail badge\n"
        "      * Expandable accordion for failed rules (sorted by severity)\n"
        "      * Evidence text highlighted in coral with optional page reference\n"
        "      * Remarks in italics\n\n"
        "  f. Dark Mode — persisted in localStorage, toggled via Header button.\n\n"
        "No bounding-box overlays or confidence scores are displayed (planned for "
        "future iteration).\""
    )

    # ══════════════════════════════════════════════════════════════
    # 13. EVALUATION METRICS
    # ══════════════════════════════════════════════════════════════
    h("13. Evaluation Metrics")
    body(
        "The planned section defines targets for OCR CER, latency, precision, recall, F1, "
        "and audit accuracy. Our project has not run formal benchmarking."
    )
    bold("Required change:")
    body(
        "Add a note: \"Formal evaluation metrics have not been established or measured. "
        "Benchmarking against a labeled dataset is planned for the next development "
        "phase.\" Alternatively, add an Appendix describing planned metrics and targets."
    )

    # ══════════════════════════════════════════════════════════════
    # 14. MODULAR ARCHITECTURE  ← REWRITE
    # ══════════════════════════════════════════════════════════════
    h("14. Modular Architecture  —  REWRITE")
    body(
        "The planned linear pipeline (OCR → Spatial → Semantic → Compliance) does not "
        "match our LangGraph-based state-graph architecture."
    )
    bold("Rewrite the diagram and description to:")
    body(
        "\"14. Modular Architecture\n\n"
        "The system is built on LangGraph, a state-graph orchestration framework. "
        "The audit pipeline consists of five graph nodes, each operating on a shared "
        "AuditState object:\n\n"
        "                  START\n"
        "                    |\n"
        "                    v\n"
        "   [1] upload_folder  ── scan directory, discover supported files\n"
        "                    |\n"
        "                    v\n"
        "   [2] summarize_documents  ── Gemini extracts type, fields, summary per doc\n"
        "                    |\n"
        "                    v\n"
        "   [3] load_checklist  ── parse checklist.md into AuditChecklist\n"
        "                    |\n"
        "                    v\n"
        "   [4] audit_agent  ── Gemini evaluates each doc against checklist\n"
        "                    |    (conditional loop — one pass per document)\n"
        "                    v\n"
        "   [5] aggregate_results  ── overall score, risk classification,\n"
        "                              prediction interval, final report\n"
        "                    |\n"
        "                    v\n"
        "                  END\n\n"
        "Key components (each independently replaceable):\n"
        "  - VLM Client (app/vlm.py): wraps Google Gemini API\n"
        "  - Checklist Parser (in app/graph.py): regex-based Markdown -> AuditRule\n"
        "  - Report Generator (backend/report_generator.py): .docx construction\n"
        "  - Backend API (backend/server.py): FastAPI with two endpoints\n"
        "  - Frontend (audit-frontend/): React SPA\""
    )

    # ══════════════════════════════════════════════════════════════
    # 15. FUTURE ROADMAP
    # ══════════════════════════════════════════════════════════════
    h("15. Future Roadmap  —  REWRITE")
    body(
        "The planned roadmap places VLM integration at Weeks 6–7. Our project already has "
        "VLM integration complete. The roadmap should be updated to reflect current reality."
    )
    bold("Rewrite to:")
    body(
        "\"15. Future Roadmap (Post-Week 5)\n\n"
        "Weeks 6–7:  ChromaDB-based Vector RAG for semantic regulatory-policy compliance\n"
        "Weeks 8–9:  Multi-document correlation (PO ↔ Invoice ↔ Ledger validation)\n"
        "Weeks 10–11: Offline fallback engine (template-based coordinate extraction)\n"
        "Week 12:    Benchmarking against labeled dataset, bounding-box overlay on previews,\n"
        "            UI refinements, project packaging and thesis documentation\""
    )

    # ══════════════════════════════════════════════════════════════
    # SUMMARY TABLE
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h("Summary of Required Changes", level=1)
    table(
        ["Section", "Change Type", "Effort"],
        [
            ["1. Document Ingestion", "Minor — add TIFF/BMP", "Low"],
            ["2. Document Rendering", "None needed", "—"],
            ["3. OCR & Layout Extraction", "MAJOR — replace with VLM description", "High"],
            ["4. Simulation / Fallback", "Delete or move to Future Work", "Low"],
            ["5. Spatial Heuristic Engine", "MAJOR — replace with VLM reasoning", "High"],
            ["6. Semantic Parsing", "Rewrite — single VLM call", "Medium"],
            ["7. Compliance Rules", "Replace examples with actual R001–R012", "Medium"],
            ["8. Compliance Scoring", "Rewrite — equal-weight, not weighted", "Medium"],
            ["9. Risk Classification", "None needed", "—"],
            ["10. Audit Trail", "Add cross-reference note", "Low"],
            ["11. Report Generation", "Rewrite — .docx not PDF", "Medium"],
            ["12. User Interface", "FULL REWRITE — React not Streamlit", "High"],
            ["13. Evaluation Metrics", "Add benchmarking status note", "Low"],
            ["14. Modular Architecture", "Rewrite — LangGraph not linear pipeline", "High"],
            ["15. Future Roadmap", "VLM already done, reorder milestones", "Medium"],
        ],
    )

    out_path = r"C:\Users\adars\OneDrive\Desktop\AUDIT_AGENT\Plan_Amendments.docx"
    doc.save(out_path)
    print(f"Saved to {out_path}")


if __name__ == "__main__":
    build()
