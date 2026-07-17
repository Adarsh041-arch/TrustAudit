import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from backend.schemas import AuditResponse


def _set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def generate_audit_docx(response: AuditResponse) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # -- Title --
    title = doc.add_heading(response.report.audit_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Elapsed: {response.elapsed_seconds:.1f}s").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # -- Overall result --
    doc.add_heading("Overall Result", level=1)
    result_label = {
        "PASS": "PASSED",
        "REVIEW REQUIRED": "REVIEW REQUIRED",
        "FAIL": "FAILED",
    }.get(response.report.overall_result, response.report.overall_result)

    p = doc.add_paragraph()
    run = p.add_run(f"{result_label}  —  {response.report.overall_score:.1f}%")
    run.bold = True
    run.font.size = Pt(16)
    if response.report.overall_result == "PASS":
        run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    elif response.report.overall_result == "FAIL":
        run.font.color.rgb = RGBColor(0x99, 0x3C, 0x1D)
    else:
        run.font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)

    pi = response.prediction_interval
    doc.add_paragraph(f"Prediction interval: {pi.lower:.1f}% — {pi.upper:.1f}%")

    doc.add_paragraph(
        f"Documents: {response.report.documents_processed} processed, "
        f"{response.report.documents_passed} passed, {response.report.documents_failed} failed"
    )
    doc.add_paragraph(response.report.summary)
    doc.add_paragraph()

    # -- Per-document sections --
    for i, dr in enumerate(response.report.document_results, 1):
        doc.add_heading(f"{i}. {dr.document_name}", level=2)

        # Preview image
        if dr.preview_base64:
            try:
                import base64
                img_bytes = base64.b64decode(dr.preview_base64)
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(2.0))
                doc.add_paragraph()
            except Exception:
                pass

        # Score + status
        p = doc.add_paragraph()
        p.add_run("Score: ").bold = True
        p.add_run(f"{dr.score:.1f}%    ")
        status_text = "PASS" if dr.passed else "FAIL"
        run = p.add_run(status_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56) if dr.passed else RGBColor(0x99, 0x3C, 0x1D)

        # Remarks
        if dr.remarks:
            doc.add_paragraph(f"Remarks: {dr.remarks}")

        # Failed rules table
        if dr.failed_rules:
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for j, text in enumerate(["Rule", "Severity", "Evidence"]):
                hdr[j].text = text
                _set_cell_shading(hdr[j], "E1F5EE")
                for paragraph in hdr[j].paragraphs:
                    for run_ in paragraph.runs:
                        run_.bold = True
                        run_.font.size = Pt(10)

            for fr in dr.failed_rules:
                row = table.add_row().cells
                row[0].text = f"{fr.rule_id}: {fr.rule_title}"
                row[1].text = fr.severity.upper()
                evidence = fr.evidence
                if fr.page_number is not None:
                    evidence += f" (p. {fr.page_number})"
                row[2].text = evidence
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run_ in paragraph.runs:
                            run_.font.size = Pt(10)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
